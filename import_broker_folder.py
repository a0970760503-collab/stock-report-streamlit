import hashlib
import re
import sqlite3
from pathlib import Path

try:
    from docx import Document
except Exception:
    Document = None

try:
    import openpyxl
except Exception:
    openpyxl = None


ROOT = Path(r"C:\Users\a0970\OneDrive\桌面\券商")
DB_PATH = Path("broker_reports.db")

BROKER_NAMES = [
    "國泰", "統一", "元大", "富邦", "兆豐", "台新", "群益", "永豐", "永豐投顧", "宏遠", "宏遠投顧",
    "華南", "合庫", "國票", "中信", "凱基", "玉山", "第一金", "第一金投顧", "上海", "台灣工銀",
    "MS", "Morgan Stanley", "GS", "Goldman Sachs", "Daiwa", "CLSA", "UBS", "Citi", "HSBC",
]

RATING_TERMS = [
    "買進", "加碼", "增持", "推薦", "區間操作", "持有", "中立", "觀望", "減碼", "賣出",
    "Buy", "Overweight", "Outperform", "Hold", "Neutral", "Sell",
]


def number(value):
    text = str(value or "").replace(",", "").strip()
    if not text or text.upper() in {"#N/A", "NA", "N/A", "NULL", "-"}:
        return None
    try:
        return float(text)
    except Exception:
        return None


def normalize_text(text):
    return re.sub(r"\s+", " ", text or "").strip()


def file_sha256(path):
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def report_file_id(path, content_hash):
    key = f"{content_hash}:{path}".encode("utf-8", errors="ignore")
    return hashlib.sha256(key).hexdigest()


def read_docx_text(path):
    if Document is None:
        return ""
    try:
        doc = Document(str(path))
        parts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text]
        for table in doc.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells if cell.text)
        return "\n".join(parts)
    except Exception:
        return ""


def read_xlsx_text(path):
    if openpyxl is None:
        return ""
    try:
        workbook = openpyxl.load_workbook(path, read_only=True, data_only=True)
        parts = []
        for sheet in workbook.worksheets:
            parts.append(f"[{sheet.title}]")
            for row in sheet.iter_rows(max_row=80, values_only=True):
                values = [str(value) for value in row if value is not None]
                if values:
                    parts.append(" ".join(values))
        return "\n".join(parts)
    except Exception:
        return ""


def extract_stock(filename):
    clean = re.sub(r"20\d{2}[-_/\.]?\d{1,2}[-_/\.]?\d{1,2}", " ", filename)
    clean = re.sub(r"20\d{2}\s*AGM", " ", clean, flags=re.IGNORECASE)
    match = re.search(r"(?<!\d)(\d{4})(?!\d)", clean)
    return match.group(1) if match else ""


def extract_date(filename, mtime):
    patterns = [
        r"(20\d{2})[-_/\.](\d{1,2})[-_/\.](\d{1,2})",
        r"(20\d{2})(\d{2})(\d{2})",
    ]
    for pattern in patterns:
        match = re.search(pattern, filename)
        if match:
            year, month, day = [int(value) for value in match.groups()]
            if 1 <= month <= 12 and 1 <= day <= 31:
                return f"{year:04d}/{month:02d}/{day:02d}"
    return mtime.strftime("%Y/%m/%d")


def extract_broker(filename):
    for broker in sorted(BROKER_NAMES, key=len, reverse=True):
        if re.search(rf"(?<![A-Za-z]){re.escape(broker)}(?![A-Za-z])", filename, re.IGNORECASE):
            return broker
    stem = Path(filename).stem
    tokens = re.split(r"[-_｜\s]+", stem)
    for token in reversed(tokens):
        token = token.strip("()（） ")
        if token and not re.fullmatch(r"\d{4,8}", token) and not re.search(r"^\d+$", token):
            return token
    return "-"


def extract_rating(text):
    source = normalize_text(text)
    for term in RATING_TERMS:
        if re.search(re.escape(term), source, re.IGNORECASE):
            return term
    return ""


def extract_target(text):
    clean = normalize_text(text).replace(",", "")
    patterns = [
        r"TP\D{0,20}([1-9]\d{1,5}(?:\.\d+)?)",
        r"目標(?:價|價格)?\D{0,16}([1-9]\d{1,5}(?:\.\d+)?)",
        r"Target\s*Price\D{0,16}([1-9]\d{1,5}(?:\.\d+)?)",
    ]
    for pattern in patterns:
        match = re.search(pattern, clean, re.IGNORECASE)
        if match:
            return number(match.group(1))

    stem = Path(text).stem
    stem = re.sub(r"20\d{2}[-_/\.](\d{1,2})[-_/\.](\d{1,2})", " ", stem)
    stem = re.sub(r"20\d{6}", " ", stem)
    stock = extract_stock(stem)
    candidates = []
    for token in re.split(r"[-_｜()\s]+", stem):
        value = number(token)
        if (
            value is not None
            and 10 <= value <= 20000
            and str(int(value)) != stock
            and not re.fullmatch(r"20\d{2}", token)
            and not re.fullmatch(r"20\d{6}", token)
        ):
            candidates.append(value)
    return candidates[-1] if candidates else None


def extract_company(filename, stock):
    stem = Path(filename).stem
    if not stock:
        return ""
    tail = stem.split(stock, 1)[-1]
    tail = re.sub(r"^[\s\-_]+", "", tail)
    tail = re.split(r"[-_｜\s]*(?:20\d{2}[-_/\.]?\d{1,2}[-_/\.]?\d{1,2}|20\d{6})", tail)[0]
    tail = re.split(r"[-_｜]", tail)[0]
    tail = re.sub(r"\([^)]*\)|（[^）]*）", "", tail).strip()
    return tail[:40]


def reason_from_text(text):
    clean = normalize_text(text)
    if not clean:
        return ""
    sentences = re.split(r"(?<=[。！？；;])\s*|\n+", clean)
    keywords = ["評等", "目標價", "上修", "下修", "營收", "毛利", "淨利", "EPS", "成長", "展望", "風險", "需求"]
    scored = []
    for index, sentence in enumerate(sentences):
        score = sum(1 for keyword in keywords if keyword.lower() in sentence.lower())
        if score:
            scored.append((score, -index, sentence.strip(" ，,。；;")))
    selected = [sentence for _, _, sentence in sorted(scored, reverse=True)[:3] if sentence]
    return " / ".join(selected)


def init_db(conn):
    conn.execute(
        """
            CREATE TABLE IF NOT EXISTS broker_files (
                file_id TEXT PRIMARY KEY,
                content_hash TEXT,
                path TEXT NOT NULL,
                relative_path TEXT NOT NULL,
                filename TEXT NOT NULL,
            extension TEXT NOT NULL,
            size INTEGER,
            modified_at TEXT
        )
        """
    )
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS broker_reports (
            file_id TEXT PRIMARY KEY,
            file_name TEXT NOT NULL,
            file_path TEXT NOT NULL,
            stock TEXT,
            company TEXT,
            broker TEXT,
            report_date TEXT,
            rating TEXT,
            target_price REAL,
            reason TEXT,
            raw_text TEXT,
            parse_status TEXT,
            imported_at TEXT DEFAULT CURRENT_TIMESTAMP
        )
        """
    )
    columns = [row[1] for row in conn.execute("PRAGMA table_info(broker_files)").fetchall()]
    if "content_hash" not in columns:
        conn.execute("ALTER TABLE broker_files ADD COLUMN content_hash TEXT")


def main():
    if not ROOT.exists():
        raise SystemExit(f"Folder not found: {ROOT}")

    files = [path for path in ROOT.rglob("*") if path.is_file()]
    with sqlite3.connect(DB_PATH) as conn:
        init_db(conn)
        conn.execute("DELETE FROM broker_files")
        conn.execute("DELETE FROM broker_reports")

        inserted = 0
        text_extracted = 0
        for path in files:
            stat = path.stat()
            rel_path = str(path.relative_to(ROOT))
            content_hash = file_sha256(path)
            file_id = report_file_id(rel_path, content_hash)
            extension = path.suffix.lower()
            mtime = stat.st_mtime
            modified_at = __import__("datetime").datetime.fromtimestamp(mtime).strftime("%Y-%m-%d %H:%M:%S")

            raw_text = ""
            parse_status = "filename_only"
            if extension == ".docx":
                raw_text = read_docx_text(path)
                parse_status = "docx_text" if raw_text else "filename_only"
            elif extension == ".xlsx":
                raw_text = read_xlsx_text(path)
                parse_status = "xlsx_text" if raw_text else "filename_only"
            elif extension == ".pdf":
                parse_status = "pdf_filename_only"

            if raw_text:
                text_extracted += 1

            filename = path.name
            source_text = f"{filename}\n{raw_text}"
            stock = extract_stock(filename)
            broker = extract_broker(filename)
            report_date = extract_date(filename, __import__("datetime").datetime.fromtimestamp(mtime))
            rating = extract_rating(source_text)
            target_price = extract_target(source_text)
            company = extract_company(filename, stock)
            reason = reason_from_text(raw_text)

            conn.execute(
                """
                INSERT INTO broker_files
                    (file_id, content_hash, path, relative_path, filename, extension, size, modified_at)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (file_id, content_hash, str(path), rel_path, filename, extension, stat.st_size, modified_at),
            )
            conn.execute(
                """
                INSERT INTO broker_reports
                    (file_id, file_name, file_path, stock, company, broker, report_date, rating,
                     target_price, reason, raw_text, parse_status, imported_at)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, CURRENT_TIMESTAMP)
                """,
                (
                    file_id,
                    filename,
                    str(path),
                    stock,
                    company,
                    broker,
                    report_date,
                    rating,
                    target_price,
                    reason,
                    raw_text[:50000],
                    parse_status,
                ),
            )
            inserted += 1

        counts = conn.execute(
            "SELECT extension, COUNT(*) FROM broker_files GROUP BY extension ORDER BY extension"
        ).fetchall()
        parsed = conn.execute(
            "SELECT parse_status, COUNT(*) FROM broker_reports GROUP BY parse_status ORDER BY parse_status"
        ).fetchall()

    print(f"folder={ROOT}")
    print(f"database={DB_PATH.resolve()}")
    print(f"files_imported={inserted}")
    print(f"text_extracted={text_extracted}")
    print("extensions=" + ", ".join(f"{ext}:{count}" for ext, count in counts))
    print("parse_status=" + ", ".join(f"{status}:{count}" for status, count in parsed))


if __name__ == "__main__":
    main()
